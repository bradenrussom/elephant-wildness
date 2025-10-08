"""
Document Processor - Core module for loading and protecting Word documents
Adapted from existing working MVP Standards Checker app
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
import re
from typing import Tuple, Optional, List


class DocumentProcessor:
    """Handles Word document loading, section identification, and content protection"""
    
    def __init__(self, file_path: str):
        """
        Initialize the document processor
        
        Args:
            file_path: Path to the .docx file
        """
        self.file_path = file_path
        self.doc = Document(file_path)
        self.page_copy_range = None
        self.disclaimer_range = None
        self._identify_sections()
    
    def _identify_sections(self):
        """
        Identify document sections using markers:
        - start_page_copy / end_page_copy
        - start_disclaimer / end_disclaimer
        """
        page_copy_start = None
        page_copy_end = None
        disclaimer_start = None
        disclaimer_end = None
        
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip().lower()
            
            # Page copy markers
            if text == 'start_page_copy':
                page_copy_start = i
            elif text == 'end_page_copy':
                page_copy_end = i
            
            # Disclaimer markers
            elif text == 'start_disclaimer':
                disclaimer_start = i
            elif text == 'end_disclaimer':
                disclaimer_end = i
        
        # Set ranges
        if page_copy_start is not None and page_copy_end is not None:
            self.page_copy_range = (page_copy_start, page_copy_end)
        
        if disclaimer_start is not None and disclaimer_end is not None:
            self.disclaimer_range = (disclaimer_start, disclaimer_end)
    
    def get_page_copy_paragraphs(self) -> List:
        """Get all paragraphs within the page copy section"""
        if self.page_copy_range:
            start, end = self.page_copy_range
            return self.doc.paragraphs[start+1:end]
        return self.doc.paragraphs
    
    def get_disclaimer_paragraphs(self) -> List:
        """Get all paragraphs within the disclaimer section"""
        if self.disclaimer_range:
            start, end = self.disclaimer_range
            return self.doc.paragraphs[start+1:end]
        return []
    
    def get_all_text(self, section: str = 'page_copy') -> str:
        """
        Get all text from specified section
        
        Args:
            section: 'page_copy', 'disclaimer', or 'all'
        
        Returns:
            Combined text from all paragraphs in section
        """
        if section == 'page_copy':
            paragraphs = self.get_page_copy_paragraphs()
        elif section == 'disclaimer':
            paragraphs = self.get_disclaimer_paragraphs()
        else:
            paragraphs = self.doc.paragraphs
        
        return '\n'.join([p.text for p in paragraphs])
    
    @staticmethod
    def protect_content(text: str) -> Tuple[str, dict]:
        """
        Protect URLs, bracketed content, and angle bracket content from modification
        
        Args:
            text: The text to protect
        
        Returns:
            Tuple of (protected_text, placeholders_dict)
        """
        placeholders = {}
        placeholder_counter = 0
        
        # Protect URLs (http://, https://, www.)
        url_pattern = r'(https?://[^\s]+|www\.[^\s]+)'
        
        def replace_url(match):
            nonlocal placeholder_counter
            placeholder = f"__URL_PLACEHOLDER_{placeholder_counter}__"
            placeholders[placeholder] = match.group(0)
            placeholder_counter += 1
            return placeholder
        
        text = re.sub(url_pattern, replace_url, text)
        
        # Protect bracketed content [like this]
        bracket_pattern = r'\[([^\]]+)\]'
        
        def replace_bracket(match):
            nonlocal placeholder_counter
            placeholder = f"__BRACKET_PLACEHOLDER_{placeholder_counter}__"
            placeholders[placeholder] = match.group(0)
            placeholder_counter += 1
            return placeholder
        
        text = re.sub(bracket_pattern, replace_bracket, text)
        
        # Protect angle bracket content <like-this>
        angle_pattern = r'<([^>]+)>'
        
        def replace_angle(match):
            nonlocal placeholder_counter
            placeholder = f"__ANGLE_PLACEHOLDER_{placeholder_counter}__"
            placeholders[placeholder] = match.group(0)
            placeholder_counter += 1
            return placeholder
        
        text = re.sub(angle_pattern, replace_angle, text)
        
        return text, placeholders
    
    @staticmethod
    def restore_content(text: str, placeholders: dict) -> str:
        """
        Restore protected content after modifications
        
        Args:
            text: Text with placeholders
            placeholders: Dictionary mapping placeholders to original content
        
        Returns:
            Text with original protected content restored
        """
        for placeholder, original in placeholders.items():
            text = text.replace(placeholder, original)
        return text
    
    def apply_text_to_paragraph(self, paragraph, new_text: str):
        """
        Apply new text to a paragraph while preserving formatting
        
        Args:
            paragraph: The paragraph object to modify
            new_text: The new text to apply
        """
        # Keep first run's formatting
        if paragraph.runs:
            first_run = paragraph.runs[0]
            
            # Clear all runs
            for run in paragraph.runs:
                run.text = ''
            
            # Set new text in first run
            first_run.text = new_text
        else:
            # No existing runs, add new one
            paragraph.add_run(new_text)
    
    def save(self, output_path: Optional[str] = None) -> str:
        """
        Save the document
        
        Args:
            output_path: Optional output path. If None, overwrites original.
        
        Returns:
            Path where document was saved
        """
        save_path = output_path or self.file_path
        self.doc.save(save_path)
        return save_path
    
    def add_analysis_section(self, analysis_text: str):
        """
        Add an analysis section to the end of the document
        
        Args:
            analysis_text: The analysis report text to append
        """
        # Add page break
        self.doc.add_page_break()
        
        # Add heading
        heading = self.doc.add_paragraph()
        heading_run = heading.add_run('Document Analysis Report')
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0, 51, 102)  # MVP blue
        
        # Add analysis text
        for line in analysis_text.split('\n'):
            if line.strip():
                para = self.doc.add_paragraph(line)
                para.style = 'Normal'
    
    def get_document_stats(self) -> dict:
        """
        Get basic document statistics
        
        Returns:
            Dictionary with word_count, paragraph_count, sentence_count
        """
        text = self.get_all_text('page_copy')
        
        # Count words
        words = text.split()
        word_count = len(words)
        
        # Count paragraphs
        paragraphs = self.get_page_copy_paragraphs()
        paragraph_count = len([p for p in paragraphs if p.text.strip()])
        
        # Count sentences (approximate)
        sentences = re.split(r'[.!?]+', text)
        sentence_count = len([s for s in sentences if s.strip()])
        
        return {
            'word_count': word_count,
            'paragraph_count': paragraph_count,
            'sentence_count': sentence_count
        }


# Example usage
if __name__ == '__main__':
    # Load document
    processor = DocumentProcessor('example.docx')
    
    # Get stats
    stats = processor.get_document_stats()
    print(f"Words: {stats['word_count']}")
    print(f"Paragraphs: {stats['paragraph_count']}")
    print(f"Sentences: {stats['sentence_count']}")
    
    # Get page copy text
    text = processor.get_all_text('page_copy')
    
    # Protect content
    protected_text, placeholders = DocumentProcessor.protect_content(text)
    
    # Do some modifications...
    # modified_text = some_function(protected_text)
    
    # Restore content
    # final_text = DocumentProcessor.restore_content(modified_text, placeholders)
    
    # Save
    # processor.save('output.docx')
